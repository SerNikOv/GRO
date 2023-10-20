import configparser

from docx import Document


def DocxProp():
    document = Document('uuuy.docx')
    core_properties = document.core_properties
    print(core_properties.DocProperty.object)

    # core_properties.author = 'Brian'
    # core_properties.author

def changeDocx(template_file_path, variables):
    output_file_path = template_file_path

    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)
