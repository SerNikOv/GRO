from docx import Document

def main():
    template_file_path = 'rrr.docx'
    output_file_path = 'result.docx'

    variables = {
        '/gro/': "302929393",
        '~grodolzhn~': '222',
        '$grofio': '333',
        "kkk": "ooo",
        "По трассе ": "302929393",
        "${EMPLOEE_PHONE}": "+972-5056000000",
        "${EMPLOEE_EMAIL}": "example@example.com",
        "${START_DATE}": "03 Jan, 2021",
        "${SALARY}": "10,000",
        "${SALARY_30}": "3,000",
        "${SALARY_70}": "7,000",
    }

    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


main()